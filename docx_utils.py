from docx import Document
from docx.shared import Inches

def s(_str, n):
    if len(_str) < n:
        return _str + (n-len(_str))*" "
    else:
#         print(n)
        return 'increase n'

def create_table(cl_mask):
    cl_v = v[cl_mask]
    argsort = np.argsort(-np.abs(cl_v))
    cl_names = np.array(topics_unique)[cl_mask]
    cl_idx = np.array(indices_unique)[cl_mask]
    cl_v_norm = cl_v / np.linalg.norm(cl_v)

    cl_names = cl_names[argsort]
    cl_idx = cl_idx[argsort]
    cl_v_norm = np.abs(cl_v_norm[argsort])
    cl_v = np.abs(cl_v[argsort])
    res = ''
    for i, (idx, name, u, u_) in enumerate(zip(cl_idx, cl_names, cl_v, cl_v_norm)):
        res += (
            f"{i + 1:<2} {s(idx, 30)} {s(name, 51)} {round(u, 3) : <5} {round(u_, 3) :<5}\n")
    return res

D = np.diag(np.sum(sim_mat, axis=1))
L = D - sim_mat


l, V = np.linalg.eigh(L)
l_mask = l > 0.1

print(len(l_mask) - l_mask.sum())

l_c = l[mask]
V_c = V[:, mask]
L_c = V_c@np.diag(l_c)@V_c.T


emb_dim = 8*8
emb = V[:, :]
emb_c = V_c[:, :emb_dim]

A = blur_array(5, size=len(emb) - 1)
document = Document()
for i in range(0, len(emb)):
    v = emb[:, i]
    argsort = np.argsort(v)
    _v = v[argsort]
    diff = A @ np.diff(_v)  # last point has no diff
    diff = np.append(diff, diff[-1])  # to account for last point

    #     # bringing the order back
    diff = diff[np.argsort(argsort)]

    d_mask = (diff > 0.001)

    pos_class = d_mask & (v > 0)
    neg_class = d_mask & (v < 0)

    #     print(topics_unique[pos_class])
    plt.figure(figsize=(3, 2))
    plt.scatter(np.arange(len(_v)), _v, c=d_mask.astype(np.int0)[argsort])
    plt.savefig('test.jpg')
    plt.close()

    document.add_heading(f'EIGENVALUE  # {i} (l = {round(l[i], 4)}) \n')
    document.add_picture('test.jpg')
    document.add_paragraph("Positive class \n" + create_table(pos_class))
    document.add_paragraph("Negative class \n" + create_table(neg_class))

document.save('demo.docx')


